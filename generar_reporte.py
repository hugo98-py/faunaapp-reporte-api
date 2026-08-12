"""
Generador de Reporte de Terreno (Word) - FaunaApp
--------------------------------------------------
MODO LOCAL: los datos se simulan con un diccionario Python (equivalente
al JSON que en el futuro vendrá desde la página Reporte de la app).
No hay servidor, no hay API: se corre directo con `python generar_reporte.py`.

Genera: portada completa + Antecedentes generales + Tabla 2 (Singularidades
ambientales registradas) + Tabla 3 (Abundancia y distribución por estación).
"""

import copy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml


# ---------------------------------------------------------------------------
# 1. DATOS SIMULADOS (esto reemplaza al JSON que enviará la página Reporte)
# ---------------------------------------------------------------------------
# NOTA: cuando pasemos a modo API, este dict será literalmente el body del
# POST que llegue desde Flutter. La forma se mantiene igual a propósito.

REPORTE_JSON_SIMULADO = {
    "metadata": {
        "nombreProyecto": "Proyecto Minero Los Bronces",
        "epoca": "Verano",
        "anio": "2025",
        "mesAnio": "Febrero 2025",
        "fechaCampana": "1 a 10 de febrero de 2025",
        "equipoProfesional": [
            {"nombre": "Javiera Soto", "profesion": "Bióloga", "jefeTerreno": True},
            {"nombre": "Matías Reyes", "profesion": "Ingeniero en RRNN", "jefeTerreno": False},
            {"nombre": "Pedro Alarcón", "profesion": "Guía de campo", "jefeTerreno": False},
        ],
        "equiposUtilizados": "Cámaras trampa, detector de ultrasonido, GPS, binoculares",
    },
    # Tabla 2: una fila por especie singular (ya resumido/agrupado en la app)
    "singularidadesTabla2": [
        {
            "clase": "Anfibios",
            "nombreCientifico": "Rhinella spinulosa",
            "nombreComun": "Sapo espinoso",
            "categoriaConservacion": "Casi amenazada (RCE)",
            "movilidad": "Baja",
            "origen": "Nativa",
        },
        {
            "clase": "Reptiles",
            "nombreCientifico": "Liolaemus nigroviridis",
            "nombreComun": "Lagartija verde negra",
            "categoriaConservacion": "-",
            "movilidad": "Baja",
            "origen": "Nativa y endémica",
        },
        {
            "clase": "Aves",
            "nombreCientifico": "Vultur gryphus",
            "nombreComun": "Cóndor andino",
            "categoriaConservacion": "Vulnerable (RCE)",
            "movilidad": "Alta",
            "origen": "Nativa",
        },
        {
            "clase": "Aves",
            "nombreCientifico": "Rhea pennata",
            "nombreComun": "Ñandú",
            "categoriaConservacion": "En peligro (RCE)",
            "movilidad": "Alta",
            "origen": "Nativa",
        },
    ],
    # Tabla 3: filas planas (clase, especie, estación, abundancia).
    # Es el mismo formato "aplanado" que ya usamos en calcularSingularidadesPorEstacion.
    "singularidadesTabla3": [
        {"clase": "Anfibios", "nombreCientifico": "Rhinella spinulosa", "nameest": "E01", "abundancia": 3},
        {"clase": "Anfibios", "nombreCientifico": "Rhinella spinulosa", "nameest": "E04", "abundancia": 1},
        {"clase": "Anfibios", "nombreCientifico": "Rhinella spinulosa", "nameest": "E06", "abundancia": 2},
        {"clase": "Reptiles", "nombreCientifico": "Liolaemus nigroviridis", "nameest": "E02", "abundancia": 5},
        {"clase": "Reptiles", "nombreCientifico": "Liolaemus nigroviridis", "nameest": "E05", "abundancia": 1},
        {"clase": "Aves", "nombreCientifico": "Vultur gryphus", "nameest": "E03", "abundancia": 4},
        {"clase": "Aves", "nombreCientifico": "Rhea pennata", "nameest": "E01", "abundancia": 1},
        {"clase": "Aves", "nombreCientifico": "Rhea pennata", "nameest": "E06", "abundancia": 1},
    ],
}


# ---------------------------------------------------------------------------
# 2. CONFIGURACIÓN / ESTILO
# ---------------------------------------------------------------------------

LOGO_PATH = Path(__file__).parent / "assets" / "logo.jpeg"
OUTPUT_PATH = Path(__file__).parent / "Reporte_Terreno_generado.docx"

COLOR_HEADER_TABLA = "FFC000"  # ambar de encabezados de tabla (Tabla 1/2/3 del template)

# Colores corporativos AMS (hex exactos entregados por el usuario)
COLOR_BARRA_VERDE = "39B74D"
COLOR_BARRA_AMBAR = "F7BC00"
COLOR_BARRA_AZUL = "5A98D1"

FONT_NAME = "Calibri"

# Tamaño de página Carta (US Letter) y márgenes, igual al template original (12240x15840 twips)
PAGE_WIDTH = Cm(21.59)
PAGE_HEIGHT = Cm(27.94)
MARGIN_TOP_BOTTOM = Cm(2.5)
MARGIN_LEFT_RIGHT = Cm(3.0)

ANCHO_BARRA = Cm(2.3)


def set_cell_shading(cell, hex_color: str):
    """Aplica color de fondo a una celda de tabla (python-docx no lo expone directo)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, align=None, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_centered_paragraph(doc, text, bold=True, size=14, color=None, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ---------------------------------------------------------------------------
# Barras de portada: se clona la ESTRUCTURA EXACTA del template original
# (grupo flotante de 3 rectángulos con relleno degradado), cambiando solo
# los colores y el alto real. Al ser una figura flotante (no una tabla), no
# participa del flujo de texto/paginación -> cero riesgo de saltos de página raros.
# ---------------------------------------------------------------------------

BARRAS_POS_H_OFFSET_EMU = 565785   # offset validado visualmente contra el original
BARRAS_ANCHO_EMU = 800100          # 2.22 cm, igual al template original

_NS_DRAWING = (
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
    'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
)


def _rect_xml(shape_id, name, off_x, off_y, ext_cx, ext_cy, gradient_stops_xml):
    return f'''<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
        xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
        <wps:cNvPr id="{shape_id}" name="{name}"/>
        <wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
        <wps:spPr bwMode="auto">
            <a:xfrm><a:off x="{off_x}" y="{off_y}"/><a:ext cx="{ext_cx}" cy="{ext_cy}"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            <a:gradFill rotWithShape="1"><a:gsLst>{gradient_stops_xml}</a:gsLst>
                <a:lin ang="5400000" scaled="1"/></a:gradFill>
            <a:ln><a:noFill/></a:ln>
        </wps:spPr>
        <wps:bodyPr rot="0" vert="horz" wrap="square" lIns="91440" tIns="45720" rIns="91440"
            bIns="45720" anchor="t" anchorCtr="0" upright="1"><a:noAutofit/></wps:bodyPr>
    </wps:wsp>'''


def _gs(pos, hex_color, shade=None):
    if shade is None:
        return f'<a:gs pos="{pos}"><a:srgbClr val="{hex_color}"/></a:gs>'
    return (f'<a:gs pos="{pos}"><a:srgbClr val="{hex_color}">'
            f'<a:gamma/><a:shade val="{shade}"/><a:invGamma/></a:srgbClr></a:gs>')


def build_barras_portada(paragraph, colores, alto_total_emu, ancho_emu=BARRAS_ANCHO_EMU):
    """Inserta el grupo flotante de 3 barras (misma estructura que el template
    original) en el párrafo dado. `colores` = [verde, ambar, azul] en hex."""
    verde, ambar, azul = colores

    # Sistema de coordenadas interno del grupo (igual al original — es solo
    # proporcional; el tamaño real lo define chExt/a:ext de más abajo).
    ch_off_x, ch_off_y = 750, 1007
    ch_ext_cx, ch_ext_cy = 975, 13660
    alto_barra_ch = 4500  # cada barra ocupa 4500 de 13660 unidades internas

    rect1 = _rect_xml(4, "Rectangle 3", 750, 1007, 975, alto_barra_ch,
                       _gs(0, verde, 78431) + _gs(100000, verde))
    rect2 = _rect_xml(5, "Rectangle 6", 750, 5594, 975, alto_barra_ch,
                       _gs(0, ambar) + _gs(50000, ambar, 78431) + _gs(100000, ambar))
    rect3 = _rect_xml(6, "Rectangle 7", 750, 10167, 975, alto_barra_ch,
                       _gs(0, azul) + _gs(100000, azul, 76078))

    xml = f'''<w:r {_NS_DRAWING}><w:drawing>
      <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"
          relativeHeight="251659264" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
        <wp:simplePos x="0" y="0"/>
        <wp:positionH relativeFrom="leftMargin"><wp:posOffset>{BARRAS_POS_H_OFFSET_EMU}</wp:posOffset></wp:positionH>
        <wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>
        <wp:extent cx="{ancho_emu}" cy="{alto_total_emu}"/>
        <wp:effectExtent l="0" t="0" r="0" b="9525"/>
        <wp:wrapNone/>
        <wp:docPr id="3" name="Group 10"/>
        <wp:cNvGraphicFramePr><a:graphicFrameLocks/></wp:cNvGraphicFramePr>
        <a:graphic>
          <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup">
            <wpg:wgp>
              <wpg:cNvGrpSpPr><a:grpSpLocks/></wpg:cNvGrpSpPr>
              <wpg:grpSpPr bwMode="auto">
                <a:xfrm><a:off x="0" y="0"/><a:ext cx="{ancho_emu}" cy="{alto_total_emu}"/>
                  <a:chOff x="{ch_off_x}" y="{ch_off_y}"/><a:chExt cx="{ch_ext_cx}" cy="{ch_ext_cy}"/></a:xfrm>
              </wpg:grpSpPr>
              {rect1}{rect2}{rect3}
            </wpg:wgp>
          </a:graphicData>
        </a:graphic>
      </wp:anchor>
    </w:drawing></w:r>'''

    run_element = parse_xml(xml)
    paragraph._p.append(run_element)


def remove_table_borders(table):
    """Deja una tabla totalmente invisible (se usa como contenedor de layout)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def set_row_height(row, height, rule=WD_ROW_HEIGHT_RULE.AT_LEAST):
    row.height = height
    row.height_rule = rule


def set_row_cant_split(row):
    """Evita que una fila de tabla se corte entre dos páginas (todo o nada)."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def merge_vertical(cells):
    """Fusiona una lista de celdas de la misma columna (celdas consecutivas de una fila cada una)."""
    base = cells[0]
    for extra in cells[1:]:
        base = base.merge(extra)
    return base


# ---------------------------------------------------------------------------
# 3. SECCIONES DEL DOCUMENTO
# ---------------------------------------------------------------------------

def build_portada(doc, metadata: dict):
    """Portada: barras de color corporativas (figura flotante, misma estructura
    que el template original) + logo/título centrados a la derecha, recreando
    tamaños y mayúsculas exactas del template."""

    contenido_alto_emu = int(PAGE_HEIGHT - MARGIN_TOP_BOTTOM * 2)
    colores = [COLOR_BARRA_VERDE, COLOR_BARRA_AMBAR, COLOR_BARRA_AZUL]

    # Párrafo ancla para la figura flotante: debe ser el primer párrafo del
    # documento (la barra se posiciona respecto a su borde superior).
    p_ancla = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    build_barras_portada(p_ancla, colores, contenido_alto_emu)

    # Indent igual al template original (ind left=426 twips) para que el
    # texto quede centrado despejado de la barra.
    indent_titulo = Cm(0.75)

    def linea_portada(texto, size_pt, space_before=0, space_after=6, bold=True):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent_titulo
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(texto)
        run.bold = bold
        run.font.size = Pt(size_pt)
        run.font.name = FONT_NAME
        return p

    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.paragraph_format.left_indent = indent_titulo
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Cm(4.5))

    linea_portada("", 20, space_before=40)
    linea_portada("REPORTE DE TERRENO", 20)
    linea_portada("LÍNEA DE BASE ECOSISTEMAS TERRESTRES", 20)
    linea_portada("FAUNA TERRESTRE", 20)
    linea_portada(metadata["nombreProyecto"], 20)
    linea_portada(f"{metadata['epoca'].upper()} - {metadata['anio']}", 20)
    linea_portada(metadata["mesAnio"].upper(), 16, space_before=170)

    doc.add_page_break()

def build_antecedentes_generales(doc, metadata: dict):
    # Título repetido tras la portada, igual al template original (16pt bold, centrado)
    add_centered_paragraph(doc, "REPORTE DE TERRENO", size=16, space_after=12)

    doc.add_heading("Antecedentes generales", level=1)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(11.0)

    # Fila 1: Fecha de campaña
    set_cell_text(table.cell(0, 0), "Fecha de campaña", bold=True)
    set_cell_text(table.cell(0, 1), metadata["fechaCampana"])

    # Fila 2: Equipo profesional (una línea por persona, jefe de terreno marcado)
    set_cell_text(table.cell(1, 0), "Equipo profesional", bold=True)
    cell_equipo = table.cell(1, 1)
    cell_equipo.text = ""
    for i, persona in enumerate(metadata["equipoProfesional"]):
        p = cell_equipo.paragraphs[0] if i == 0 else cell_equipo.add_paragraph()
        etiqueta = " (jefe/a de terreno)" if persona.get("jefeTerreno") else ""
        run = p.add_run(f"{persona['nombre']}{etiqueta}, {persona['profesion']}")
        run.font.size = Pt(10)
        run.font.name = FONT_NAME

    # Fila 3: Equipos utilizados
    set_cell_text(table.cell(2, 0), "Equipos utilizados", bold=True)
    set_cell_text(table.cell(2, 1), metadata["equiposUtilizados"])

    doc.add_paragraph()


def build_tabla2_singularidades(doc, filas: list[dict]):
    doc.add_heading("Resultados", level=1)
    doc.add_heading("Singularidades ambientales", level=2)

    p_caption = doc.add_paragraph()
    run = p_caption.add_run("Tabla 2. Singularidades ambientales registradas")
    run.bold = True
    run.font.size = Pt(10)

    headers = ["Clase", "Nombre científico", "Nombre común", "Categoría de conservación", "Movilidad", "Origen"]
    col_widths = [2.3, 3.5, 3.0, 3.5, 1.8, 2.4]

    # Agrupar filas por clase para poder fusionar la primera columna
    clases_orden = []
    por_clase: dict[str, list[dict]] = {}
    for f in filas:
        clase = f["clase"]
        if clase not in por_clase:
            por_clase[clase] = []
            clases_orden.append(clase)
        por_clase[clase].append(f)

    total_filas = 1 + sum(len(v) for v in por_clase.values())  # +1 header
    table = doc.add_table(rows=total_filas, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(col_widths):
        table.columns[i].width = Cm(w)

    for c, h in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_text(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, COLOR_HEADER_TABLA)

    row_idx = 1
    for clase in clases_orden:
        especies = por_clase[clase]
        start_row = row_idx
        for especie in especies:
            set_cell_text(table.cell(row_idx, 1), especie["nombreCientifico"])
            set_cell_text(table.cell(row_idx, 2), especie["nombreComun"])
            set_cell_text(table.cell(row_idx, 3), especie["categoriaConservacion"], align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(table.cell(row_idx, 4), especie["movilidad"], align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(table.cell(row_idx, 5), especie["origen"], align=WD_ALIGN_PARAGRAPH.CENTER)
            row_idx += 1
        end_row = row_idx - 1
        # Fusionar celdas de "Clase" para ese grupo
        cells_clase = [table.cell(r, 0) for r in range(start_row, end_row + 1)]
        merged = merge_vertical(cells_clase)
        set_cell_text(merged, clase, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_fuente = doc.add_paragraph()
    run = p_fuente.add_run("Fuente: AMS Consultores, 2025.")
    run.font.size = Pt(9)
    run.italic = True
    doc.add_paragraph()


def build_tabla3_por_estacion(doc, filas: list[dict]):
    p_caption = doc.add_paragraph()
    run = p_caption.add_run("Tabla 3. Abundancia y distribución de singularidades ambientales por estación de muestreo")
    run.bold = True
    run.font.size = Pt(10)

    # Estaciones únicas presentes en los datos (no todas las de la campaña, solo las con registro)
    estaciones = sorted({f["nameest"] for f in filas})

    # Agrupar clase -> especie -> {nameest: abundancia}
    estructura: dict[str, dict[str, dict[str, int]]] = {}
    clases_orden = []
    for f in filas:
        clase, especie, est, cant = f["clase"], f["nombreCientifico"], f["nameest"], f["abundancia"]
        if clase not in estructura:
            estructura[clase] = {}
            clases_orden.append(clase)
        estructura[clase].setdefault(especie, {})
        estructura[clase][especie][est] = estructura[clase][especie].get(est, 0) + cant

    n_cols = 2 + len(estaciones)
    n_filas_especies = sum(len(especies) for especies in estructura.values())
    table = doc.add_table(rows=2 + n_filas_especies, cols=n_cols)  # 2 filas de encabezado
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [2.3, 3.5] + [1.6] * len(estaciones)
    for i, w in enumerate(col_widths):
        table.columns[i].width = Cm(w)

    # Encabezado fila 1: "Clase" / "Nombre científico" (merge vertical con fila 2) + "Estación de muestreo" (merge horizontal)
    merge_vertical([table.cell(0, 0), table.cell(1, 0)])
    set_cell_text(table.cell(0, 0), "Clase", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, 0), COLOR_HEADER_TABLA)

    merge_vertical([table.cell(0, 1), table.cell(1, 1)])
    set_cell_text(table.cell(0, 1), "Nombre científico", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, 1), COLOR_HEADER_TABLA)

    if estaciones:
        header_est_cells = [table.cell(0, 2 + i) for i in range(len(estaciones))]
        merged_est = merge_vertical(header_est_cells) if len(header_est_cells) > 1 else header_est_cells[0]
        set_cell_text(merged_est, "Estación de muestreo", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(merged_est, COLOR_HEADER_TABLA)

    for i, est in enumerate(estaciones):
        cell = table.cell(1, 2 + i)
        set_cell_text(cell, est, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, COLOR_HEADER_TABLA)

    row_idx = 2
    for clase in clases_orden:
        especies = estructura[clase]
        start_row = row_idx
        for especie, valores in especies.items():
            set_cell_text(table.cell(row_idx, 1), especie)
            for i, est in enumerate(estaciones):
                abundancia = valores.get(est, 0)
                texto = str(abundancia) if abundancia else "-"
                set_cell_text(table.cell(row_idx, 2 + i), texto, align=WD_ALIGN_PARAGRAPH.CENTER)
            row_idx += 1
        end_row = row_idx - 1
        cells_clase = [table.cell(r, 0) for r in range(start_row, end_row + 1)]
        merged = merge_vertical(cells_clase)
        set_cell_text(merged, clase, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    p_fuente = doc.add_paragraph()
    run = p_fuente.add_run("Fuente: AMS Consultores, 2025.")
    run.font.size = Pt(9)
    run.italic = True


# ---------------------------------------------------------------------------
# 4. ORQUESTADOR
# ---------------------------------------------------------------------------

def generar_reporte(data: dict, output_path: Path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP_BOTTOM
    section.bottom_margin = MARGIN_TOP_BOTTOM
    section.left_margin = MARGIN_LEFT_RIGHT
    section.right_margin = MARGIN_LEFT_RIGHT

    build_portada(doc, data["metadata"])
    build_antecedentes_generales(doc, data["metadata"])
    build_tabla2_singularidades(doc, data["singularidadesTabla2"])
    build_tabla3_por_estacion(doc, data["singularidadesTabla3"])

    doc.save(str(output_path))
    print(f"Documento generado: {output_path}")


if __name__ == "__main__":
    generar_reporte(REPORTE_JSON_SIMULADO, OUTPUT_PATH)
